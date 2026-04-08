"""
connectors/resume_parser.py
═══════════════════════════════════════════════════════
Parses PDF and DOCX resumes into structured text
that can be fed into the Scorer Agent.

Supports:
  - PDF  → via pymupdf (fitz)
  - DOCX → via python-docx
  - TXT  → plain text fallback

Flow:
  Resume file → extract raw text → clean → structure
  → returns ResumeParseResult ready for Applicant model
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import fitz  # pymupdf
from docx import Document

from utils.logger import logger


# ─────────────────────────────────────────────────────
#  Result Model
# ─────────────────────────────────────────────────────

@dataclass
class ResumeParseResult:
    """
    Structured output from resume parsing.
    Fed directly into the Applicant model fields.
    """
    raw_text:           str                     = ""
    full_name:          Optional[str]           = None
    email:              Optional[str]           = None
    phone:              Optional[str]           = None
    location:           Optional[str]           = None
    github_url:         Optional[str]           = None
    linkedin_url:       Optional[str]           = None
    portfolio_url:      Optional[str]           = None
    skills:             list[str]               = field(default_factory=list)
    education:          Optional[str]           = None
    experience_entries: list[dict]              = field(default_factory=list)
    total_pages:        int                     = 0
    file_type:          str                     = ""
    parse_success:      bool                    = False
    error_message:      Optional[str]           = None

    def is_valid(self) -> bool:
        """Return True if minimum useful data was extracted."""
        return self.parse_success and len(self.raw_text.strip()) > 100


# ─────────────────────────────────────────────────────
#  Text Extraction
# ─────────────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> tuple[str, int]:
    """
    Extract raw text from PDF bytes using pymupdf.
    Returns (text, page_count).
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = []
        for page in doc:
            pages.append(page.get_text("text"))
        doc.close()
        return "\n".join(pages), len(pages)
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise


def extract_text_from_docx(file_bytes: bytes) -> tuple[str, int]:
    """
    Extract raw text from DOCX bytes using python-docx.
    Returns (text, paragraph_count).
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs), len(paragraphs)
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise


def extract_text_from_txt(file_bytes: bytes) -> tuple[str, int]:
    """Extract text from plain .txt file."""
    text = file_bytes.decode("utf-8", errors="ignore")
    lines = [l for l in text.splitlines() if l.strip()]
    return text, len(lines)


# ─────────────────────────────────────────────────────
#  Text Cleaning
# ─────────────────────────────────────────────────────

def clean_text(text: str) -> str:
    """
    Clean extracted resume text.
    Removes noise while preserving structure.
    """
    # Normalize whitespace
    text = re.sub(r'\r\n', '\n', text)
    text = re.sub(r'\r', '\n', text)

    # Remove excessive blank lines (keep max 2)
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Remove weird unicode characters
    text = re.sub(r'[^\x00-\x7F\u00C0-\u024F\u1E00-\u1EFF]', ' ', text)

    # Remove lines that are just symbols/noise
    lines = text.splitlines()
    cleaned = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            cleaned.append("")
            continue
        # Skip lines that are just dots, dashes, underscores
        if re.match(r'^[\.\-_\*\|=\s]{3,}$', stripped):
            continue
        cleaned.append(stripped)

    return "\n".join(cleaned).strip()


# ─────────────────────────────────────────────────────
#  Field Extraction (regex-based)
# ─────────────────────────────────────────────────────

def extract_email(text: str) -> Optional[str]:
    """Extract first email address found in text."""
    pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    match = re.search(pattern, text)
    return match.group(0).lower() if match else None


def extract_phone(text: str) -> Optional[str]:
    """Extract first phone number found in text."""
    patterns = [
        r'\+?91[-\s]?\d{10}',           # Indian mobile with country code
        r'\+?\d{1,3}[-\s]?\d{10}',      # International
        r'\d{10}',                       # Plain 10-digit
        r'\(\d{3}\)\s?\d{3}[-\s]\d{4}', # US format
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(0).strip()
    return None


def extract_github_url(text: str) -> Optional[str]:
    """Extract GitHub profile URL."""
    pattern = r'https?://(?:www\.)?github\.com/[\w\-]+'
    match = re.search(pattern, text, re.IGNORECASE)
    return match.group(0) if match else None


def extract_linkedin_url(text: str) -> Optional[str]:
    """Extract LinkedIn profile URL."""
    pattern = r'https?://(?:www\.)?linkedin\.com/in/[\w\-]+'
    match = re.search(pattern, text, re.IGNORECASE)
    return match.group(0) if match else None


def extract_portfolio_url(text: str) -> Optional[str]:
    """
    Extract portfolio/personal website URL.
    Excludes GitHub and LinkedIn.
    """
    pattern = r'https?://(?!(?:www\.)?(?:github|linkedin|twitter|instagram)\.com)[\w\-\.]+\.[a-z]{2,}(?:/[\w\-\./?=&]*)?'
    matches = re.findall(pattern, text, re.IGNORECASE)
    # Return first non-social URL
    for url in matches:
        if not any(skip in url for skip in ['github', 'linkedin', 'twitter', 'instagram']):
            return url
    return None


def extract_skills(text: str) -> list[str]:
    """
    Extract tech skills from resume text.
    Uses a curated list of common tech skills.
    """
    # Common tech skills to look for
    TECH_SKILLS = [
        # Languages
        "Python", "Java", "JavaScript", "TypeScript", "C++", "C#", "Go",
        "Rust", "Kotlin", "Swift", "R", "Scala", "PHP", "Ruby",
        # Web
        "React", "Next.js", "Vue", "Angular", "Node.js", "FastAPI",
        "Django", "Flask", "Spring Boot", "Express", "HTML", "CSS",
        # Data / ML
        "TensorFlow", "PyTorch", "Scikit-learn", "Pandas", "NumPy",
        "Keras", "Hugging Face", "LangChain", "OpenCV", "NLTK", "spaCy",
        "Matplotlib", "Seaborn", "Plotly",
        # Databases
        "PostgreSQL", "MySQL", "MongoDB", "Redis", "Elasticsearch",
        "Supabase", "Firebase", "SQLite", "Cassandra", "DynamoDB",
        # Cloud / DevOps
        "AWS", "GCP", "Azure", "Docker", "Kubernetes", "Terraform",
        "CI/CD", "GitHub Actions", "Jenkins", "Ansible",
        # Tools
        "Git", "Linux", "REST API", "GraphQL", "Kafka", "RabbitMQ",
        "Airflow", "Spark", "Hadoop", "dbt",
        # AI/LLM
        "LLM", "RAG", "Groq", "OpenAI", "Anthropic", "LangGraph",
        "Vector Database", "Embeddings", "Fine-tuning",
    ]

    found = []
    text_lower = text.lower()
    for skill in TECH_SKILLS:
        # Match whole word, case-insensitive
        pattern = r'\b' + re.escape(skill.lower()) + r'\b'
        if re.search(pattern, text_lower):
            found.append(skill)

    return found


def extract_education(text: str) -> Optional[str]:
    """
    Extract education section from resume.
    Returns first education entry found.
    """
    edu_keywords = [
        "b.tech", "b.e.", "btech", "b.sc", "bsc", "m.tech", "mtech",
        "m.sc", "msc", "mba", "phd", "bachelor", "master", "degree",
        "university", "college", "institute", "iit", "nit", "bits",
    ]
    lines = text.splitlines()
    edu_lines = []
    capture = False

    for i, line in enumerate(lines):
        line_lower = line.lower()
        # Start capturing near education keywords
        if any(kw in line_lower for kw in edu_keywords):
            capture = True
        if capture:
            edu_lines.append(line.strip())
            if len(edu_lines) >= 4:  # capture up to 4 lines
                break

    return " | ".join(l for l in edu_lines if l) if edu_lines else None


def extract_name_from_top(text: str) -> Optional[str]:
    """
    Heuristic: first non-empty line is usually the candidate's name.
    Filter out obvious non-names.
    """
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    if not lines:
        return None

    first_line = lines[0]

    # Skip if it looks like a URL, email, or header
    skip_patterns = ['http', '@', 'resume', 'cv', 'curriculum']
    if any(p in first_line.lower() for p in skip_patterns):
        return None

    # Skip if too long (probably not a name)
    if len(first_line) > 60:
        return None

    # Skip if it's all caps and very short (likely a section header)
    if first_line.isupper() and len(first_line) < 5:
        return None

    return first_line


# ─────────────────────────────────────────────────────
#  Main Parser
# ─────────────────────────────────────────────────────

class ResumeParser:
    """
    Main resume parser class.
    Supports PDF, DOCX, and TXT files.

    Usage:
        parser = ResumeParser()
        result = parser.parse(file_bytes=b"...", file_type="pdf")
    """

    SUPPORTED_TYPES = {"pdf", "docx", "doc", "txt"}

    def parse(
        self,
        file_bytes: bytes,
        file_type: str,
        filename: str = "",
    ) -> ResumeParseResult:
        """
        Parse a resume file and return structured result.

        Args:
            file_bytes: Raw file content as bytes
            file_type:  File extension (pdf, docx, txt)
            filename:   Original filename for logging

        Returns:
            ResumeParseResult with extracted fields
        """
        result = ResumeParseResult(file_type=file_type.lower())
        file_type = file_type.lower().strip(".")

        if file_type not in self.SUPPORTED_TYPES:
            result.error_message = f"Unsupported file type: {file_type}"
            logger.warning(f"Resume parse skipped — unsupported type: {file_type}")
            return result

        try:
            logger.info(f"Parsing resume | Type: {file_type} | File: {filename}")

            # ── Step 1: Extract raw text ──────────
            if file_type == "pdf":
                raw_text, pages = extract_text_from_pdf(file_bytes)
                result.total_pages = pages
            elif file_type in ("docx", "doc"):
                raw_text, pages = extract_text_from_docx(file_bytes)
                result.total_pages = pages
            else:  # txt
                raw_text, pages = extract_text_from_txt(file_bytes)
                result.total_pages = 1

            # ── Step 2: Clean text ────────────────
            cleaned = clean_text(raw_text)
            result.raw_text = cleaned

            # ── Step 3: Extract fields ────────────
            result.email         = extract_email(cleaned)
            result.phone         = extract_phone(cleaned)
            result.github_url    = extract_github_url(cleaned)
            result.linkedin_url  = extract_linkedin_url(cleaned)
            result.portfolio_url = extract_portfolio_url(cleaned)
            result.skills        = extract_skills(cleaned)
            result.education     = extract_education(cleaned)
            result.full_name     = extract_name_from_top(cleaned)

            result.parse_success = True

            logger.info(
                f"Resume parsed | Name: {result.full_name} | "
                f"Email: {result.email} | Skills: {len(result.skills)} | "
                f"Pages: {result.total_pages}"
            )

        except Exception as e:
            result.error_message = str(e)
            result.parse_success = False
            logger.error(f"Resume parse failed for {filename}: {e}")

        return result

    def parse_from_path(self, file_path: str | Path) -> ResumeParseResult:
        """
        Parse resume from a local file path.
        Useful for bulk processing and testing.
        """
        path = Path(file_path)
        if not path.exists():
            result = ResumeParseResult()
            result.error_message = f"File not found: {file_path}"
            return result

        file_bytes = path.read_bytes()
        return self.parse(
            file_bytes=file_bytes,
            file_type=path.suffix,
            filename=path.name,
        )


# ─────────────────────────────────────────────────────
#  Global Instance
# ─────────────────────────────────────────────────────

resume_parser = ResumeParser()